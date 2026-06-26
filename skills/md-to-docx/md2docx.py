#!/usr/bin/env python3
"""md2docx: Markdown -> branded Word .docx via Pandoc + a reference template.

All options are optional except --reference. Per file:
  pandoc (against the reference doc) -> optionally repoint tables to a branded table
  style -> optionally strip Heading auto-numbering / set margins (done on the reference)
  -> optionally prepend a cover built from the reference's first page, with the
  Title/Subtitle/Author/Date placeholders filled.

KEY TRICK for table column widths: Pandoc reads the *relative number of dashes* in a
Markdown table's separator row as the column widths. So to widen a column, give it more
dashes (alignment colons are preserved):
    | Element | What it is                              | Demand |
    |---------|-----------------------------------------|--------|
This lives in the Markdown source, not here, so it travels with the document.

Requires: pandoc on PATH;  pip install python-docx docxcompose
"""
import argparse, os, re, shutil, subprocess, sys, tempfile, zipfile
from docx import Document
from docx.oxml.ns import qn
try:
    from docxcompose.composer import Composer
except ImportError:
    Composer = None


def rewrite_part(path, part, fn):
    """Apply fn(text)->text to one part inside a .docx, in place."""
    with zipfile.ZipFile(path) as z:
        names = z.namelist(); data = {n: z.read(n) for n in names}
    data[part] = fn(data[part].decode('utf-8')).encode('utf-8')
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, data[n])


def first_h1(md):
    with open(md, encoding='utf-8') as f:
        for line in f:
            if line.startswith('# '):
                return line[2:].strip()
    return os.path.splitext(os.path.basename(md))[0]


def prep_reference(template, dest, strip_numbers, margins):
    shutil.copy(template, dest)
    if strip_numbers:                       # remove auto-numbering from Heading 1-9, keep outline levels
        def fix(xml):
            def nonum(m):
                return re.sub(r'(<w:numId w:val=")\d+("\s*/>)', r'\g<1>0\g<2>', m.group(0))
            return re.sub(r'<w:style [^>]*w:styleId="Heading[1-9]"[^>]*>.*?</w:style>', nonum, xml, flags=re.S)
        rewrite_part(dest, 'word/styles.xml', fix)
    if margins:                             # set page margins (twips); 720 = 0.5in
        def fix(xml):
            def narrow(m):
                tag = m.group(0)
                for a in ('top', 'right', 'bottom', 'left'):
                    tag = re.sub(rf'w:{a}="\d+"', f'w:{a}="{margins}"', tag)
                return tag
            return re.sub(r'<w:pgMar\b[^>]*>', narrow, xml)
        rewrite_part(dest, 'word/document.xml', fix)


def convert(md, reference, out, table_style):
    subprocess.run(['pandoc', md, '-o', out, '--reference-doc', reference], check=True)
    if table_style:                         # Pandoc tags tables "Table"; point them at the branded style
        rewrite_part(out, 'word/document.xml',
                     lambda x: re.sub(r'<w:tblStyle w:val="Table"\s*/>',
                                      f'<w:tblStyle w:val="{table_style}"/>', x))


def build_cover(reference, dest):
    """A cover is the reference's first page: strip the body paragraphs/tables, keep the
    cover drawings/content-controls and the section properties."""
    shutil.copy(reference, dest)
    d = Document(dest); body = d.element.body
    sect = body.find(qn('w:sectPr'))
    for p in list(d.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(d.tables):
        t._element.getparent().remove(t._element)
    if sect is not None and body.find(qn('w:sectPr')) is None:
        body.append(sect)
    d.save(dest)


def fill_cover(path, mapping):
    """Replace placeholder text inside cover text boxes. Matches whole-paragraph text so
    it works even when the placeholder is split across runs."""
    d = Document(path)
    for tx in d.element.iter(qn('w:txbxContent')):
        for p in tx.iter(qn('w:p')):
            ts = p.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in ts).strip()
            if text in mapping and ts:
                ts[0].text = mapping[text]
                for t in ts[1:]:
                    t.text = ''
    d.save(path)


def merge(cover, body, out):
    if Composer is None:
        sys.exit("docxcompose not installed:  pip install docxcompose")
    master = Document(cover); master.add_page_break()
    Composer(master).append(Document(body)); master.save(out)


def main():
    ap = argparse.ArgumentParser(description="Markdown -> branded Word .docx via Pandoc + reference template")
    ap.add_argument('files', nargs='+', help='input .md file(s)')
    ap.add_argument('-r', '--reference', required=True, help='reference .docx (styles, margins, headers/footers, cover)')
    ap.add_argument('-o', '--outdir', default='.', help='output directory (default: current dir)')
    ap.add_argument('--table-style', help='styleId to apply to tables, e.g. ListTable4-Accent1 (must exist in the reference)')
    ap.add_argument('--no-heading-numbers', action='store_true', help='strip Heading auto-numbering (use when the Markdown already has manual section numbers)')
    ap.add_argument('--margins', type=int, metavar='TWIPS', help='page margins in twips (720 = 0.5in narrow)')
    ap.add_argument('--cover', action='store_true', help='prepend a cover built from the reference doc first page')
    ap.add_argument('--title', help="cover Title (default: each file's first '# ' heading)")
    ap.add_argument('--subtitle', default='', help='cover Subtitle')
    ap.add_argument('--author', default='', help='cover Author')
    ap.add_argument('--date', default='', help='cover Date')
    ap.add_argument('--placeholders', default='Document Title,Document Subtitle,Author Name,Date',
                    help='comma list of the cover placeholder texts to match, in order: title,subtitle,author,date')
    a = ap.parse_args()

    if not shutil.which('pandoc'):
        sys.exit("pandoc not found on PATH")
    os.makedirs(a.outdir, exist_ok=True)
    ph = [s.strip() for s in a.placeholders.split(',')]

    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, 'reference.docx')
        prep_reference(a.reference, ref, a.no_heading_numbers, a.margins)
        base_cover = os.path.join(tmp, 'cover.docx')
        if a.cover:
            build_cover(ref, base_cover)
        for md in a.files:
            name = os.path.splitext(os.path.basename(md))[0] + '.docx'
            out = os.path.join(a.outdir, name)
            body = os.path.join(tmp, 'body-' + name)
            convert(md, ref, body, a.table_style)
            if a.cover:
                cov = os.path.join(tmp, 'cov-' + name)
                shutil.copy(base_cover, cov)
                fill_cover(cov, dict(zip(ph, [a.title or first_h1(md), a.subtitle, a.author, a.date])))
                merge(cov, body, out)
            else:
                shutil.copy(body, out)
            print('wrote', out)


if __name__ == '__main__':
    main()
